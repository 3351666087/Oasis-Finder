from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

from .config import get_settings
from .services import build_network_graph, load_dashboard_snapshot, load_network_data
from .db import create_app_engine
import pandas as pd
from sqlalchemy import text


def _engine():
    return create_app_engine()


def _build_risk_chart(output_path: Path) -> None:
    snapshot = load_dashboard_snapshot()
    chart_df = snapshot.risk_distribution.groupby("tier_level", as_index=False)["avg_score"].mean()
    plt.figure(figsize=(7.8, 4.6))
    plt.bar(chart_df["tier_level"], chart_df["avg_score"], color=["#205072", "#329D9C", "#56C596", "#7BE495", "#B5E48C", "#D9ED92"])
    plt.title("Average Risk Score by Network Tier")
    plt.ylabel("Risk Score")
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()


def _build_demand_chart(output_path: Path) -> None:
    snapshot = load_dashboard_snapshot()
    plt.figure(figsize=(8.4, 4.6))
    plt.plot(snapshot.demand_trend["business_date"], snapshot.demand_trend["units_sold"], color="#1F6F8B", linewidth=2.2)
    plt.fill_between(snapshot.demand_trend["business_date"], snapshot.demand_trend["units_sold"], color="#99C1DE", alpha=0.35)
    plt.title("Network Demand Trend - Last 60 Days")
    plt.ylabel("Units Sold")
    plt.xticks(rotation=20)
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()


def _build_network_chart(output_path: Path) -> None:
    graph = build_network_graph()
    plt.figure(figsize=(11.5, 6.4))
    tier_order = {"L3": 0, "L2": 1, "L1": 2, "CORE": 3, "DOWNSTREAM": 4, "SERVICE": 5}
    node_groups = {}
    for node, data in graph.nodes(data=True):
        node_groups.setdefault(data["tier"], []).append(node)

    positions = {}
    for tier, nodes in node_groups.items():
        y_values = list(range(len(nodes)))
        x_value = tier_order.get(tier, 6)
        for idx, node in enumerate(nodes):
            positions[node] = (x_value, -y_values[idx])

    colors = {
        "L1": "#2A9D8F",
        "L2": "#1D3557",
        "L3": "#E76F51",
        "CORE": "#264653",
        "DOWNSTREAM": "#F4A261",
        "SERVICE": "#8D99AE",
    }
    nx = __import__("networkx")
    node_colors = [colors.get(graph.nodes[node]["tier"], "#6C757D") for node in graph.nodes]
    nx.draw_networkx_edges(graph, positions, alpha=0.35, edge_color="#4F5D75", arrows=True, arrowsize=9)
    nx.draw_networkx_nodes(graph, positions, node_color=node_colors, node_size=260, edgecolors="#F7F7F7", linewidths=0.5)
    labels = {node: node for node in list(graph.nodes)[:36]}
    nx.draw_networkx_labels(graph, positions, labels=labels, font_size=6, font_color="#1B1B1B")
    plt.title("Mesh Supply Chain Topology with L1/L2/L3 Tiers")
    plt.axis("off")
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()


def build_presubmission_report(output_path: Path) -> Path:
    settings = get_settings()
    settings.report_asset_root.mkdir(parents=True, exist_ok=True)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    risk_chart = settings.report_asset_root / "risk_by_tier.png"
    demand_chart = settings.report_asset_root / "demand_trend.png"
    network_chart = settings.report_asset_root / "network_topology.png"

    _build_risk_chart(risk_chart)
    _build_demand_chart(demand_chart)
    _build_network_chart(network_chart)

    snapshot = load_dashboard_snapshot()
    stats = pd.read_sql(
        text(
            """
            SELECT
                (SELECT COUNT(*) FROM organizations) AS organizations,
                (SELECT COUNT(*) FROM facilities) AS facilities,
                (SELECT COUNT(*) FROM supply_edges) AS supply_edges,
                (SELECT COUNT(*) FROM supplier_lots) AS supplier_lots,
                (SELECT COUNT(*) FROM product_batches) AS product_batches,
                (SELECT COUNT(*) FROM shipments) AS shipments,
                (SELECT COUNT(*) FROM demand_history) AS demand_rows
            """
        ),
        _engine(),
    ).iloc[0]

    document = Document()
    document.add_heading("INF101TC Semester 2: Pre-Submission 1", 0)
    document.add_paragraph("Group Number: 12")
    document.add_paragraph("Project Title: Mesh Supply Chain Intelligence System")
    document.add_paragraph("Team Deliverable Support: Codex-assisted engineering implementation")

    document.add_heading("Section 1: The Problem Definition", level=1)
    document.add_paragraph(
        "The business problem is the trust and traceability gap inside complex fresh-food supply chains, where multiple upstream tiers, cold-chain execution, and quality variation make incident response too slow. "
        "The system engineering challenge is to turn a fragmented supplier network into a tier-aware digital mesh that can trace batches, forecast demand, and simulate disruption recovery in near real time."
    )

    document.add_heading("Section 2: Project Scope", level=1)
    document.add_paragraph(
        "This submission delivers a Python repository backed by MySQL and a PySide6 desktop front end. "
        f"The current prototype models {int(stats['organizations'])} organisations, {int(stats['facilities'])} facilities, {int(stats['supply_edges'])} active supply links, "
        f"{int(stats['supplier_lots'])} upstream lots, {int(stats['product_batches'])} finished-goods batches, {int(stats['shipments'])} shipments, and {int(stats['demand_rows'])} demand observations."
    )
    document.add_paragraph(
        "The semester scope includes L1/L2/L3 supplier visualisation, batch traceability, risk scoring, demand forecasting, and scenario-based recovery planning. "
        "The future scope is to connect live IoT telemetry, supplier portals, and recall workflows."
    )

    document.add_heading("Section 3: Engineering Requirements", level=1)
    document.add_paragraph(
        "Functional requirements: The application must map the full network, display L1/L2/L3 suppliers, trace finished batches to upstream lots, show shipment and inspection history, forecast product demand, and simulate node disruption recovery."
    )
    document.add_paragraph(
        "Non-functional requirements: Core screens should respond within two seconds on local hardware, maintain consistent tier labelling, and support operational decisions with quantitative metrics such as risk score, fill-rate recovery, and reorder point."
    )

    document.add_heading("Section 4: The Framework / Process", level=1)
    document.add_paragraph(
        "The team process used six engineering phases: requirement extraction from the original brief, supply-network domain modelling, MySQL schema design, synthetic dataset generation, analytics model training, and PySide6 interface validation. "
        "Instead of building a static mock-up, the system was developed as a working digital twin with seeded operational data."
    )

    document.add_heading("Section 5: The Selected Approach", level=1)
    document.add_paragraph(
        "A Python code repository was selected because it allows direct integration between MySQL data engineering, machine learning analytics, simulation logic, and a deployable desktop interface. "
        "Compared with slideware or a purely conceptual design, this approach demonstrates executable engineering value and leaves a reusable platform for later expansion."
    )

    document.add_heading("Section 6: Tools Justification", level=1)
    document.add_paragraph(
        "MySQL was selected for relational traceability and supplier-network storage, PySide6 for a rich desktop operations interface, and open-source Python analytics libraries for forecasting and risk modelling. "
        "This stack balances realism, maintainability, and speed of implementation while supporting future integration into enterprise systems."
    )

    document.add_heading("Section 7: Progress on Semester Deliverables", level=1)
    document.add_paragraph(
        f"Current dashboard snapshot: {snapshot.kpis['supplier_facilities']} supplier facilities, {snapshot.kpis['active_links']} active links, {snapshot.kpis['total_batches']} product batches, "
        f"{snapshot.kpis['on_time_rate']}% on-time rate, and {snapshot.kpis['active_alerts']} active alerts."
    )

    document.add_picture(str(network_chart), width=Inches(6.4))
    document.add_paragraph("Caption: The network topology visualises the mesh structure and explicitly highlights L1, L2, and L3 supplier tiers connected to core plants and downstream fulfilment hubs.")

    document.add_picture(str(risk_chart), width=Inches(6.2))
    document.add_paragraph("Caption: The risk distribution chart demonstrates that the system does not only store nodes, but also scores operational exposure by tier to support mitigation decisions.")

    document.add_picture(str(demand_chart), width=Inches(6.2))
    document.add_paragraph("Caption: The demand trend figure demonstrates the forecasting dataset and provides evidence that the system tracks temporal network demand instead of static records only.")

    document.save(output_path)
    return output_path

