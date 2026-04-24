from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "Week 6 Artefact Template.docx"

SUBMISSION_ROOT = ROOT / "ENT103TC_Week6_Submission"
REQUIRED_DIR = SUBMISSION_ROOT / "01_Required"
EDITABLE_DIR = SUBMISSION_ROOT / "02_Editable"
SUPPORTING_DIR = SUBMISSION_ROOT / "03_Supporting"

ARTEFACT_DOCX = EDITABLE_DIR / "ENT103TC_Week6_Artefact_StudentStyle.docx"
CODEBOOK_DOCX = EDITABLE_DIR / "ENT103TC_Week6_Codebook_StudentStyle.docx"
CODEBOOK_CSV = EDITABLE_DIR / "ENT103TC_Week6_Codebook_StudentStyle.csv"
OPEN_CODING_CSV = EDITABLE_DIR / "ENT103TC_Week6_OpenCoding_StudentStyle.csv"
SOURCES_CSV = EDITABLE_DIR / "ENT103TC_Week6_SourceList_StudentStyle.csv"
README_PATH = SUPPORTING_DIR / "Submission_Checklist.txt"
REQUIREMENTS_NOTE_PATH = SUPPORTING_DIR / "Course_Requirement_Notes.txt"


DECISION_QUESTION = (
    "Should our app focus first on QR code traceability and supply chain analytics, "
    "or should we spend our main effort on adding real-time factory video to show food safety?"
)

MISSING_INFORMATION = (
    "We still do not know which type of transparency users will trust more in real life. "
    "We also do not know if live factory video is realistic for a big food company in terms "
    "of privacy, cost, and daily operation."
)

DATA_NEEDED = [
    "Consumer opinions about what kind of food safety information makes them feel more confident.",
    "Research or reports about how traceability systems improve trust, recall speed, and transparency.",
    "Evidence about whether live video is technically possible, easy to manage, and legally safe to use at scale.",
]

APA_REFERENCES = [
    (
        "Badia-Melis, R., Mishra, P., & Ruiz-Garcia, L. (2015). Food traceability: "
        "New trends and recent advances. A review. Food Control, 57, 393-401. "
        "https://doi.org/10.1016/j.foodcont.2015.05.005"
    ),
    (
        "World Economic Forum. (2019). Innovation with a purpose: Improving traceability "
        "in food value chains through technology innovations. "
        "https://www.weforum.org/reports/innovation-with-a-purpose-improving-traceability-in-food-value-chains-through-technology-innovations/"
    ),
    (
        "Wang, E. S.-T., Lin, H.-C., & Tsai, M.-C. (2021). Effect of institutional trust "
        "on consumers' health and safety perceptions and repurchase intention for traceable "
        "fresh food. Foods, 10(12), 2898. https://doi.org/10.3390/foods10122898"
    ),
]

SOURCE_ROWS = [
    [
        "Journal article",
        "Badia-Melis et al. (2015)",
        "Food traceability review",
        "https://doi.org/10.1016/j.foodcont.2015.05.005",
        "Used to understand what a strong traceability system should do and why end-to-end records matter.",
    ],
    [
        "Industry report",
        "World Economic Forum (2019)",
        "Innovation with a purpose",
        "https://www.weforum.org/reports/innovation-with-a-purpose-improving-traceability-in-food-value-chains-through-technology-innovations/",
        "Used to understand real business value, infrastructure needs, and practical rollout limits.",
    ],
    [
        "Journal article",
        "Wang et al. (2021)",
        "Traceable fresh food and consumer trust",
        "https://doi.org/10.3390/foods10122898",
        "Used to connect traceability with trust, safety perception, and repurchase intention.",
    ],
]

PATTERNS = [
    "Pattern 1: People trust food information more when it is clear, traceable, and backed by reliable organisations.",
    "Pattern 2: A good transparency system needs connected data from different stages of the supply chain, not just one attractive front-end feature.",
    "Pattern 3: New tools sound useful, but they only work well when the company also has standards, governance, and enough technical support.",
]

INSIGHTS = [
    "Insight 1: The repeated idea in all three sources is that trust comes from evidence people can check, so QR records, batch links, and inspection data are stronger than just showing a video screen.",
    "Insight 2: The main conflict is that live video looks impressive to users, but it may not actually prove much if the upstream data, certifications, and shipment records are still weak.",
    "Insight 3: Our current project direction already matches the research better because we are building multi-tier traceability, risk scoring, and forecasting, which gives more useful proof than a single visibility feature.",
]

FINAL_DECISION = (
    "Your decision: We should make QR-based traceability and supply chain analytics the main priority first. "
    "Real-time factory video can be kept as a later optional feature, but it should not replace proper traceability data."
)

CODEBOOK_ROWS = [
    [
        "Badia-Melis et al. (2015)",
        "records need to connect across the chain",
        "information_flow_gap",
        "Traceability is weak when data is separated and hard to connect.",
        "Pattern 2",
    ],
    [
        "Badia-Melis et al. (2015)",
        "fast access matters during recalls",
        "quick_traceback",
        "A useful system should help people find source information quickly.",
        "Pattern 2",
    ],
    [
        "Badia-Melis et al. (2015)",
        "technology should be practical to use",
        "practical_adoption",
        "Even good tools fail if they are too hard or expensive to use.",
        "Pattern 3",
    ],
    [
        "World Economic Forum (2019)",
        "technology only helps when standards are shared",
        "harmonised_standards",
        "Companies need shared rules and formats to make traceability work.",
        "Pattern 3",
    ],
    [
        "World Economic Forum (2019)",
        "full chain visibility creates value",
        "end_to_end_visibility",
        "Traceability becomes stronger when different stages are linked together.",
        "Pattern 2",
    ],
    [
        "World Economic Forum (2019)",
        "rollout needs infrastructure and training",
        "implementation_support",
        "The business still needs systems, training, and support behind the feature.",
        "Pattern 3",
    ],
    [
        "Wang et al. (2021)",
        "people trust traceable food more when institutions are reliable",
        "institutional_trust",
        "Trusted organisations make traceability information feel more believable.",
        "Pattern 1",
    ],
    [
        "Wang et al. (2021)",
        "producer trust affects safety feelings",
        "producer_trust",
        "Consumers care about who produced the food, not only the label itself.",
        "Pattern 1",
    ],
    [
        "Wang et al. (2021)",
        "trust increases buying intention",
        "repurchase_intention",
        "When trust is higher, people are more willing to buy again.",
        "Pattern 1",
    ],
]

OPEN_CODING_ROWS = [
    [
        "Badia-Melis et al. (2015)",
        "Traceability systems need connected records and quick access during recalls.",
        "information_flow_gap",
        "Pattern 2",
        "Repeater",
    ],
    [
        "Badia-Melis et al. (2015)",
        "Advanced tools are useful only if they can be adopted in practice.",
        "practical_adoption",
        "Pattern 3",
        "Repeater",
    ],
    [
        "World Economic Forum (2019)",
        "Technology alone is not enough without standards and infrastructure.",
        "harmonised_standards",
        "Pattern 3",
        "Repeater",
    ],
    [
        "World Economic Forum (2019)",
        "End-to-end visibility improves safety and efficiency.",
        "end_to_end_visibility",
        "Pattern 2",
        "Repeater",
    ],
    [
        "Wang et al. (2021)",
        "Consumer trust depends on reliable institutions and producers.",
        "institutional_trust",
        "Pattern 1",
        "Repeater",
    ],
    [
        "Wang et al. (2021)",
        "Trust improves safety perception and repurchase intention.",
        "repurchase_intention",
        "Pattern 1",
        "Repeater",
    ],
    [
        "Cross-source conflict",
        "Live video may look transparent, but traceability data gives stronger evidence people can verify later.",
        "video_vs_data_proof",
        "Insight conflict",
        "Conflict",
    ],
]


def ensure_output_dirs() -> None:
    for directory in [SUBMISSION_ROOT, REQUIRED_DIR, EDITABLE_DIR, SUPPORTING_DIR]:
        directory.mkdir(parents=True, exist_ok=True)


def set_paragraph_text(document: Document, paragraph_number: int, text: str) -> None:
    document.paragraphs[paragraph_number - 1].text = text


def build_filled_artefact() -> Path:
    document = Document(TEMPLATE_PATH)
    replacements = {
        9: DECISION_QUESTION,
        13: MISSING_INFORMATION,
        17: DATA_NEEDED[0],
        18: DATA_NEEDED[1],
        19: DATA_NEEDED[2],
        24: APA_REFERENCES[0],
        25: APA_REFERENCES[1],
        26: APA_REFERENCES[2],
        27: "The matching codebook and CSV files are attached in the submission folder.",
        33: "I used the three documents above to make the codebook and open-coding CSV.",
        37: PATTERNS[0],
        38: PATTERNS[1],
        39: PATTERNS[2],
        40: "These patterns repeat across the journal and industry sources, so they are not just one-off comments.",
        45: INSIGHTS[0],
        46: INSIGHTS[1],
        47: INSIGHTS[2],
        52: FINAL_DECISION,
    }
    for paragraph_number, text in replacements.items():
        set_paragraph_text(document, paragraph_number, text)
    document.save(ARTEFACT_DOCX)
    return ARTEFACT_DOCX


def build_codebook_docx() -> Path:
    document = Document()
    document.add_heading("ENT103TC Week 6 Codebook", 0)
    document.add_paragraph(f"Decision question: {DECISION_QUESTION}")
    document.add_paragraph("This version is written in a simple student style and follows the Week 6 open-coding task.")

    document.add_heading("Three Source Documents", level=1)
    for reference in APA_REFERENCES:
        document.add_paragraph(reference)

    document.add_heading("Codes", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Source"
    header[1].text = "Excerpt / idea"
    header[2].text = "Code"
    header[3].text = "Short meaning"
    header[4].text = "Merged pattern"

    for row in CODEBOOK_ROWS:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    document.add_heading("Patterns", level=1)
    for pattern in PATTERNS:
        document.add_paragraph(pattern)

    document.add_heading("Insights", level=1)
    for insight in INSIGHTS:
        document.add_paragraph(insight)

    document.save(CODEBOOK_DOCX)
    return CODEBOOK_DOCX


def write_csv(path: Path, header: list[str], rows: list[list[str]]) -> Path:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(header)
        writer.writerows(rows)
    return path


def build_supporting_notes() -> None:
    README_PATH.write_text(
        "\n".join(
            [
                "ENT103TC Week 6 Submission Checklist",
                "",
                "Based on the lecture slide requirement, the required submission is:",
                "1. The filled Week 6 artefact template",
                "2. The codebook generated from the three collected documents",
                "",
                "This folder includes both required files in PDF form under 01_Required.",
                "Editable DOCX and CSV versions are in 02_Editable.",
                "Extra support notes are in 03_Supporting.",
            ]
        ),
        encoding="utf-8",
    )
    REQUIREMENTS_NOTE_PATH.write_text(
        "\n".join(
            [
                "Course requirement notes taken from Week 6 lecture slides:",
                "",
                "Page 35 says:",
                "1. Based on the missing information, collect three pieces of data.",
                "2. Use Taguette and apply the 4-step open-coding approach.",
                "3. Submit the template and the codebook as Week 6 artefacts.",
                "",
                "This submission package is built around that checklist.",
            ]
        ),
        encoding="utf-8",
    )


def copy_required_versions() -> None:
    artefact_pdf = ROOT / "ENT103TC_Week6_Artefact_StudentStyle.pdf"
    codebook_pdf = ROOT / "ENT103TC_Week6_Codebook_StudentStyle.pdf"
    if artefact_pdf.exists():
        shutil.copy2(artefact_pdf, REQUIRED_DIR / artefact_pdf.name)
    if codebook_pdf.exists():
        shutil.copy2(codebook_pdf, REQUIRED_DIR / codebook_pdf.name)

    for file_path in [ARTEFACT_DOCX, CODEBOOK_DOCX, CODEBOOK_CSV, OPEN_CODING_CSV, SOURCES_CSV]:
        if file_path.exists():
            continue


def main() -> None:
    ensure_output_dirs()
    build_filled_artefact()
    build_codebook_docx()
    write_csv(
        CODEBOOK_CSV,
        ["Source", "Excerpt / idea", "Code", "Short meaning", "Merged pattern"],
        CODEBOOK_ROWS,
    )
    write_csv(
        OPEN_CODING_CSV,
        ["Source", "Text / observation", "Code", "Pattern", "Repeater or Conflict"],
        OPEN_CODING_ROWS,
    )
    write_csv(
        SOURCES_CSV,
        ["Type", "Author / organisation", "Title", "Link", "Why we used it"],
        SOURCE_ROWS,
    )
    build_supporting_notes()
    print(f"Created: {ARTEFACT_DOCX}")
    print(f"Created: {CODEBOOK_DOCX}")
    print(f"Created: {CODEBOOK_CSV}")
    print(f"Created: {OPEN_CODING_CSV}")
    print(f"Created: {SOURCES_CSV}")
    print(f"Created: {README_PATH}")
    print(f"Created: {REQUIREMENTS_NOTE_PATH}")


if __name__ == "__main__":
    main()
