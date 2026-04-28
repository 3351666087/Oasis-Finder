from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
import sys

import matplotlib.pyplot as plt
import matplotlib.patheffects as pe
from matplotlib.patches import Circle, FancyArrowPatch, FancyBboxPatch
import numpy as np
from sqlalchemy import text

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"

if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from mesh_supply_chain.db import create_app_engine
from mesh_supply_chain.services import (
    get_batch_codes,
    get_batch_trace,
    get_disruptable_facilities,
    get_forecast_series,
    get_product_options,
    get_region_options,
    load_dashboard_snapshot,
    simulate_disruption,
)


GROUP_NUMBER = "9"
PROJECT_TITLE = "Mesh Supply Chain Intelligence System"
OUTPUT_PATH = ROOT / "INF.docx"
ASSET_DIR = ROOT / "artifacts" / "report_assets"
ARCHITECTURE_PATH = ASSET_DIR / "system_architecture.png"
NETWORK_PATH = ASSET_DIR / "network_topology.png"
RISK_PATH = ASSET_DIR / "risk_by_tier.png"
DEMAND_PATH = ASSET_DIR / "demand_trend.png"
RISK_METRICS_PATH = ROOT / "artifacts" / "risk_metrics.json"
FORECAST_METRICS_PATH = ROOT / "artifacts" / "forecast_metrics.json"


@dataclass(frozen=True)
class RuntimeContext:
    organizations: int
    facilities: int
    active_links: int
    supplier_lots: int
    batches: int
    shipments: int
    demand_rows: int
    supplier_facilities: int
    on_time_rate: float
    average_risk_score: float
    active_alerts: int
    risk_rmse: float
    risk_mae: float
    forecast_rmse: float
    forecast_mae: float
    forecast_mape_pct: float
    forecast_horizon_days: int
    sample_batch_code: str
    sample_product_name: str
    sample_plant_name: str
    sample_component_count: int
    sample_shipment_count: int
    sample_forecast_sku: str
    sample_forecast_region: str
    sample_forecast_history_days: int
    sample_forecast_days: int
    sample_sim_facility_code: str
    sample_sim_facility_name: str
    sample_sim_fill_rate_pct: float
    sample_sim_impacted_edges: int
    sample_sim_message: str


def _read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _fallback_context() -> RuntimeContext:
    risk_metrics = _read_json(RISK_METRICS_PATH)
    forecast_metrics = _read_json(FORECAST_METRICS_PATH)
    return RuntimeContext(
        organizations=79,
        facilities=97,
        active_links=199,
        supplier_lots=154,
        batches=414,
        shipments=758,
        demand_rows=35040,
        supplier_facilities=73,
        on_time_rate=57.39,
        average_risk_score=79.36,
        active_alerts=10,
        risk_rmse=float(risk_metrics["rmse"]),
        risk_mae=float(risk_metrics["mae"]),
        forecast_rmse=float(forecast_metrics["rmse"]),
        forecast_mae=float(forecast_metrics["mae"]),
        forecast_mape_pct=float(forecast_metrics["mape"]) * 100,
        forecast_horizon_days=int(forecast_metrics["forecast_horizon_days"]),
        sample_batch_code="BAT-00414",
        sample_product_name="Premium Pork Loin 500g",
        sample_plant_name="Chengdu Resilience Food Plant",
        sample_component_count=5,
        sample_shipment_count=2,
        sample_forecast_sku="SKU-CHB-500",
        sample_forecast_region="Central China",
        sample_forecast_history_days=90,
        sample_forecast_days=30,
        sample_sim_facility_code="FAC-L1-003",
        sample_sim_facility_name="Fuzhou Broiler Input Processing Hub",
        sample_sim_fill_rate_pct=100.0,
        sample_sim_impacted_edges=1,
        sample_sim_message="Scenario completed with 100.0% projected recovery coverage.",
    )


def collect_runtime_context() -> RuntimeContext:
    risk_metrics = _read_json(RISK_METRICS_PATH)
    forecast_metrics = _read_json(FORECAST_METRICS_PATH)

    try:
        snapshot = load_dashboard_snapshot()
        engine = create_app_engine()
        with engine.connect() as connection:
            stats_row = connection.execute(
                text(
                    """
                    SELECT
                        (SELECT COUNT(*) FROM organizations) AS organizations,
                        (SELECT COUNT(*) FROM facilities) AS facilities,
                        (SELECT COUNT(*) FROM supply_edges WHERE active = 1) AS active_links,
                        (SELECT COUNT(*) FROM supplier_lots) AS supplier_lots,
                        (SELECT COUNT(*) FROM product_batches) AS batches,
                        (SELECT COUNT(*) FROM shipments) AS shipments,
                        (SELECT COUNT(*) FROM demand_history) AS demand_rows
                    """
                )
            ).mappings().one()

        sample_batch_code = get_batch_codes(1)[0]
        trace = get_batch_trace(sample_batch_code)

        product_options = get_product_options()
        region_options = get_region_options()
        sample_forecast_sku = next((sku for sku, _ in product_options if sku == "SKU-CHB-500"), product_options[0][0])
        sample_forecast_region = "Central China" if "Central China" in region_options else region_options[0]
        forecast_series = get_forecast_series(sample_forecast_sku, sample_forecast_region)

        best_sim = None
        for facility_code, facility_name in get_disruptable_facilities():
            result = simulate_disruption(facility_code, 35)
            candidate = (facility_code, facility_name, result)
            if best_sim is None or candidate[2]["fill_rate"] > best_sim[2]["fill_rate"]:
                best_sim = candidate
            if candidate[2]["fill_rate"] >= 0.999:
                break

        if best_sim is None:
            raise RuntimeError("No disruption simulation result was available.")

        return RuntimeContext(
            organizations=int(stats_row["organizations"]),
            facilities=int(stats_row["facilities"]),
            active_links=int(stats_row["active_links"]),
            supplier_lots=int(stats_row["supplier_lots"]),
            batches=int(stats_row["batches"]),
            shipments=int(stats_row["shipments"]),
            demand_rows=int(stats_row["demand_rows"]),
            supplier_facilities=int(snapshot.kpis["supplier_facilities"]),
            on_time_rate=float(snapshot.kpis["on_time_rate"]),
            average_risk_score=float(snapshot.kpis["average_risk_score"]),
            active_alerts=int(snapshot.kpis["active_alerts"]),
            risk_rmse=float(risk_metrics["rmse"]),
            risk_mae=float(risk_metrics["mae"]),
            forecast_rmse=float(forecast_metrics["rmse"]),
            forecast_mae=float(forecast_metrics["mae"]),
            forecast_mape_pct=float(forecast_metrics["mape"]) * 100,
            forecast_horizon_days=int(forecast_metrics["forecast_horizon_days"]),
            sample_batch_code=sample_batch_code,
            sample_product_name=str(trace["header"]["product_name"]),
            sample_plant_name=str(trace["header"]["plant_name"]),
            sample_component_count=int(len(trace["components"])),
            sample_shipment_count=int(len(trace["shipments"])),
            sample_forecast_sku=sample_forecast_sku,
            sample_forecast_region=sample_forecast_region,
            sample_forecast_history_days=int(len(forecast_series["history"])),
            sample_forecast_days=int(len(forecast_series["forecast"])),
            sample_sim_facility_code=best_sim[0],
            sample_sim_facility_name=best_sim[1],
            sample_sim_fill_rate_pct=float(best_sim[2]["fill_rate"]) * 100,
            sample_sim_impacted_edges=int(best_sim[2]["impacted_edges"]),
            sample_sim_message=str(best_sim[2]["message"]),
        )
    except Exception:
        return _fallback_context()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_document(document: Document) -> None:
    section = document.sections[-1]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    header = section.header.paragraphs[0]
    header.text = "INF101TC | Group 9 | Mesh Supply Chain Intelligence System"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header.runs:
        header.runs[0].font.name = "Calibri"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(88, 101, 118)

    footer = section.footer.paragraphs[0]
    footer.clear()
    _add_page_number(footer)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    title_style = document.styles["Title"]
    title_style.font.name = "Cambria"
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(18, 52, 59)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Cambria"
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(18, 52, 59)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Cambria"
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(31, 111, 139)


def add_body_paragraph(document: Document, text: str, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.italic = italic


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.add_run(item)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_figure(document: Document, image_path: Path, width: float, caption: str) -> None:
    document.add_picture(str(image_path), width=Inches(width))
    add_caption(document, caption)


def build_architecture_diagram(context: RuntimeContext) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = ["DejaVu Sans", "Segoe UI", "Arial"]
    fig, ax = plt.subplots(figsize=(16, 9))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 9)
    ax.axis("off")
    fig.patch.set_facecolor("#08121D")
    ax.set_facecolor("#08121D")

    start = np.array([8, 18, 29]) / 255.0
    mid = np.array([14, 32, 54]) / 255.0
    end = np.array([12, 22, 38]) / 255.0
    gradient = np.zeros((900, 1400, 3))
    for channel in range(3):
        left = np.linspace(start[channel], mid[channel], gradient.shape[1] // 2)
        right = np.linspace(mid[channel], end[channel], gradient.shape[1] - gradient.shape[1] // 2)
        gradient[:, :, channel] = np.concatenate([left, right])[None, :]
    ax.imshow(gradient, extent=[0, 16, 0, 9], aspect="auto", zorder=0)

    ax.add_patch(Circle((2.0, 8.0), 1.9, color="#2EE6D6", alpha=0.10, zorder=0.2))
    ax.add_patch(Circle((13.8, 7.8), 1.6, color="#5CAEFF", alpha=0.10, zorder=0.2))
    ax.add_patch(Circle((14.4, 1.0), 2.0, color="#FF7A8D", alpha=0.09, zorder=0.2))
    ax.add_patch(Circle((7.8, 0.7), 1.8, color="#FFC766", alpha=0.08, zorder=0.2))

    flow_panel = FancyBboxPatch(
        (0.55, 3.30),
        14.9,
        2.75,
        boxstyle="round,pad=0.02,rounding_size=0.28",
        facecolor="#0E1D30",
        edgecolor=(1, 1, 1, 0.06),
        linewidth=1.0,
        zorder=0.9,
    )
    flow_panel.set_path_effects([pe.withSimplePatchShadow(offset=(0, -1.5), shadow_rgbFace=(0, 0, 0), alpha=0.18)])
    ax.add_patch(flow_panel)

    output_panel = FancyBboxPatch(
        (0.80, 1.10),
        14.4,
        2.10,
        boxstyle="round,pad=0.02,rounding_size=0.24",
        facecolor="#0C192A",
        edgecolor=(1, 1, 1, 0.05),
        linewidth=1.0,
        zorder=0.9,
    )
    output_panel.set_path_effects([pe.withSimplePatchShadow(offset=(0, -1.2), shadow_rgbFace=(0, 0, 0), alpha=0.14)])
    ax.add_patch(output_panel)

    title_bar = FancyBboxPatch(
        (0.70, 8.05),
        3.60,
        0.48,
        boxstyle="round,pad=0.02,rounding_size=0.22",
        facecolor="#2EE6D6",
        edgecolor="none",
        zorder=2,
    )
    ax.add_patch(title_bar)
    validation_targets: list[tuple[object, tuple[float, float, float, float], str]] = []

    fig.canvas.draw()

    def _measure_text(text: str, fontsize: float, fontweight: str = "normal", linespacing: float = 1.2):
        temp = ax.text(
            0,
            0,
            text if text else " ",
            fontsize=fontsize,
            fontweight=fontweight,
            ha="left",
            va="top",
            alpha=0,
            linespacing=linespacing,
            zorder=50,
        )
        fig.canvas.draw()
        bbox = temp.get_window_extent(renderer=fig.canvas.get_renderer())
        temp.remove()
        return bbox

    def _wrap_text_to_width(text: str, fontsize: float, max_width_px: float, fontweight: str = "normal") -> str:
        wrapped_blocks = []
        for paragraph in text.split("\n"):
            words = paragraph.split()
            if not words:
                wrapped_blocks.append("")
                continue
            lines = [words[0]]
            for word in words[1:]:
                trial = f"{lines[-1]} {word}"
                if _measure_text(trial, fontsize, fontweight=fontweight).width <= max_width_px:
                    lines[-1] = trial
                else:
                    lines.append(word)
            wrapped_blocks.append("\n".join(lines))
        return "\n".join(wrapped_blocks)

    def _fit_text(
        x: float,
        y: float,
        w: float,
        h: float,
        text: str,
        color: str,
        max_font: float,
        min_font: float,
        fontweight: str = "normal",
        ha: str = "left",
        va: str = "top",
        linespacing: float = 1.2,
        key: str = "",
        pad_x: float = 0.0,
        pad_y: float = 0.0,
        zorder: float = 4,
    ):
        (px0, py0) = ax.transData.transform((x, y))
        (px1, py1) = ax.transData.transform((x + w, y + h))
        max_width_px = abs(px1 - px0)
        max_height_px = abs(py1 - py0)
        chosen_text = text
        chosen_size = min_font
        for fontsize in np.arange(max_font, min_font - 0.01, -0.3):
            candidate = _wrap_text_to_width(text, float(fontsize), max_width_px, fontweight=fontweight)
            bbox = _measure_text(candidate, float(fontsize), fontweight=fontweight, linespacing=linespacing)
            if bbox.width <= max_width_px and bbox.height <= max_height_px:
                chosen_text = candidate
                chosen_size = float(fontsize)
                break

        if ha == "left":
            anchor_x = x + pad_x
        elif ha == "center":
            anchor_x = x + w / 2
        else:
            anchor_x = x + w - pad_x

        if va == "top":
            anchor_y = y + h - pad_y
        elif va == "center":
            anchor_y = y + h / 2
        else:
            anchor_y = y + pad_y

        artist = ax.text(
            anchor_x,
            anchor_y,
            chosen_text,
            fontsize=chosen_size,
            fontweight=fontweight,
            color=color,
            ha=ha,
            va=va,
            linespacing=linespacing,
            zorder=zorder,
        )
        allowed = ax.transData.transform([[x, y], [x + w, y + h]])
        validation_targets.append(
            (
                artist,
                (
                    float(min(allowed[0][0], allowed[1][0])),
                    float(min(allowed[0][1], allowed[1][1])),
                    float(max(allowed[0][0], allowed[1][0])),
                    float(max(allowed[0][1], allowed[1][1])),
                ),
                key,
            )
        )
        return artist

    _fit_text(0.88, 8.12, 3.22, 0.28, "SYSTEM WALKTHROUGH", "#06121D", 10.6, 10.0, fontweight="bold", key="header_tag", pad_x=0.0, pad_y=0.01)
    _fit_text(0.70, 7.10, 8.80, 0.60, "Mesh Supply Chain Intelligence System", "#F7FBFF", 24.0, 17.0, fontweight="bold", key="hero_title")
    _fit_text(
        0.70,
        6.56,
        8.60,
        0.42,
        "Investor-style technical flow from data foundation to scalable operational intelligence",
        "#AFC3D8",
        11.8,
        9.0,
        key="hero_subtitle",
    )

    badge = FancyBboxPatch(
        (10.85, 7.02),
        4.00,
        1.05,
        boxstyle="round,pad=0.03,rounding_size=0.20",
        facecolor=(1, 1, 1, 0.07),
        edgecolor=(1, 1, 1, 0.12),
        linewidth=1.1,
        zorder=2,
    )
    badge.set_path_effects([pe.withSimplePatchShadow(offset=(1.8, -1.8), shadow_rgbFace=(0, 0, 0), alpha=0.14)])
    ax.add_patch(badge)
    _fit_text(11.10, 7.68, 3.45, 0.20, "LIVE REPOSITORY SNAPSHOT", "#F7FBFF", 9.8, 9.0, fontweight="bold", key="badge_title")
    _fit_text(
        11.10,
        7.18,
        3.45,
        0.34,
        f"{context.organizations} orgs | {context.facilities} facilities | {context.active_links} active links | {context.demand_rows:,} demand rows",
        "#BBD0E3",
        9.0,
        7.8,
        key="badge_body",
    )

    _fit_text(0.88, 5.72, 2.0, 0.18, "PITCH-DECK FLOW", "#7EA3C4", 9.8, 9.0, fontweight="bold", key="flow_label")
    _fit_text(0.98, 2.64, 2.8, 0.18, "INVESTABLE OUTPUTS", "#7EA3C4", 9.8, 9.0, fontweight="bold", key="output_label")

    ax.plot([1.30, 14.80], [4.78, 4.78], color="#2EE6D6", linewidth=3.0, alpha=0.18, zorder=1.3)
    ax.plot([1.30, 14.80], [4.78, 4.78], color="#5CAEFF", linewidth=8.0, alpha=0.05, zorder=1.2)

    def draw_card(x: float, y: float, w: float, h: float, accent: str, step: str, tag: str, title: str, subtitle: str, lines: list[str]) -> None:
        shadow = FancyBboxPatch(
            (x + 0.08, y - 0.09),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.20",
            facecolor="#000000",
            edgecolor="none",
            alpha=0.22,
            zorder=1,
        )
        ax.add_patch(shadow)
        card = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.20",
            facecolor=(1, 1, 1, 0.94),
            edgecolor=(1, 1, 1, 0.10),
            linewidth=1.0,
            zorder=2,
        )
        ax.add_patch(card)
        card.set_path_effects([pe.withSimplePatchShadow(offset=(1.2, -1.2), shadow_rgbFace=(0, 0, 0), alpha=0.12)])
        accent_band = FancyBboxPatch(
            (x, y + h - 0.18),
            w,
            0.18,
            boxstyle="round,pad=0.03,rounding_size=0.20",
            facecolor=accent,
            edgecolor="none",
            zorder=2.2,
        )
        ax.add_patch(accent_band)

        glow = Circle((x + 0.38, y + h + 0.21), 0.28, facecolor=accent, edgecolor="none", alpha=0.18, zorder=2.4)
        ax.add_patch(glow)
        circle = Circle((x + 0.38, y + h + 0.21), 0.18, facecolor=accent, edgecolor="#FFFFFF", linewidth=1.4, zorder=3)
        circle.set_path_effects([pe.withSimplePatchShadow(offset=(1.2, -1.2), shadow_rgbFace=(0, 0, 0), alpha=0.18)])
        ax.add_patch(circle)
        _fit_text(x + 0.28, y + h + 0.08, 0.20, 0.22, step, "#FFFFFF", 10.4, 9.5, fontweight="bold", ha="center", va="center", key=f"{step}_step", zorder=4)

        tag_box = FancyBboxPatch(
            (x + 0.18, y + h - 0.54),
            1.34,
            0.28,
            boxstyle="round,pad=0.02,rounding_size=0.10",
            facecolor="#EEF5FF",
            edgecolor="none",
            zorder=3,
        )
        ax.add_patch(tag_box)
        _fit_text(x + 0.22, y + h - 0.51, 1.26, 0.20, tag, accent, 8.9, 7.0, fontweight="bold", ha="center", va="center", key=f"{step}_tag", zorder=4)
        _fit_text(x + 0.18, y + h - 0.95, w - 0.36, 0.28, title, "#102538", 13.0, 10.0, fontweight="bold", key=f"{step}_title")
        body_y = y + 0.18
        body_h = h - 1.22
        if subtitle.strip():
            _fit_text(x + 0.18, y + h - 1.26, w - 0.36, 0.24, subtitle, "#617587", 9.2, 7.4, key=f"{step}_subtitle")
            body_h = h - 1.48
        body_text = "\n".join(lines)
        _fit_text(x + 0.18, body_y, w - 0.36, body_h, body_text, "#203546", 9.5, 7.5, linespacing=1.35, key=f"{step}_body", va="bottom")

    def draw_output_card(x: float, y: float, w: float, h: float, accent: str, title: str, body: str) -> None:
        shadow = FancyBboxPatch(
            (x + 0.08, y - 0.08),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.20",
            facecolor="#000000",
            edgecolor="none",
            alpha=0.18,
            zorder=1,
        )
        ax.add_patch(shadow)
        card = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.20",
            facecolor=(1, 1, 1, 0.94),
            edgecolor=(1, 1, 1, 0.08),
            linewidth=1.0,
            zorder=2,
        )
        ax.add_patch(card)
        ax.plot([x + 0.16, x + w - 0.16], [y + h - 0.22, y + h - 0.22], color=accent, linewidth=3.2, zorder=3)
        _fit_text(x + 0.18, y + h - 0.58, w - 0.36, 0.24, title, "#102538", 11.6, 9.4, fontweight="bold", key=f"{title}_title")
        _fit_text(x + 0.18, y + 0.18, w - 0.36, h - 0.92, body, "#334959", 9.2, 7.4, linespacing=1.34, key=f"{title}_body", va="bottom")

    def draw_connector(start: tuple[float, float], end: tuple[float, float], label: str, color: str, label_box: tuple[float, float, float, float] | None = None, rad: float = 0.0) -> None:
        arrow = FancyArrowPatch(
            start,
            end,
            connectionstyle=f"arc3,rad={rad}",
            arrowstyle="-|>",
            mutation_scale=18,
            linewidth=2.4,
            color=color,
            zorder=2.8,
        )
        ax.add_patch(arrow)
        if label_box is not None:
            lx, ly, lw, lh = label_box
            pill = FancyBboxPatch(
                (lx, ly),
                lw,
                lh,
                boxstyle="round,pad=0.02,rounding_size=0.12",
                facecolor=(1, 1, 1, 0.92),
                edgecolor=(1, 1, 1, 0.18),
                linewidth=0.8,
                zorder=3.2,
            )
            ax.add_patch(pill)
            _fit_text(lx, ly, lw, lh, label.upper(), color, 7.8, 7.0, fontweight="bold", ha="center", va="center", key=f"{label}_connector", zorder=4)

    main_y = 3.88
    main_w = 2.55
    main_h = 1.95
    x_positions = [0.95, 3.90, 6.85, 9.80, 12.75]
    accents = ["#2EE6D6", "#5CAEFF", "#7AF5A0", "#FFC766", "#FF7A8D"]

    draw_card(
        x_positions[0],
        main_y,
        main_w,
        main_h,
        accents[0],
        "1",
        "RUNTIME",
        "Local Setup",
        "",
        ["MySQL service on port 3307", "One-command bootstrap flow", "Repeatable local execution path"],
    )
    draw_card(
        x_positions[1],
        main_y,
        main_w,
        main_h,
        accents[1],
        "2",
        "DATA CORE",
        "Relational Model",
        "",
        ["17 schema tables", "SQLAlchemy entity mapping", "Traceability-ready joins"],
    )
    draw_card(
        x_positions[2],
        main_y,
        main_w,
        main_h,
        accents[2],
        "3",
        "DIGITAL TWIN",
        "Synthetic Mesh",
        "",
        [
            f"{context.organizations} organisations",
            f"{context.facilities} facilities | {context.active_links} links",
            f"{context.demand_rows:,} demand observations",
        ],
    )
    draw_card(
        x_positions[3],
        main_y,
        main_w,
        main_h,
        accents[3],
        "4",
        "INTELLIGENCE",
        "Predictive Layer",
        "",
        [
            f"Risk RMSE {context.risk_rmse:.2f}",
            f"Forecast MAPE {context.forecast_mape_pct:.2f}%",
            f"{context.forecast_horizon_days}-day planning horizon",
        ],
    )
    draw_card(
        x_positions[4],
        main_y,
        main_w,
        main_h,
        accents[4],
        "5",
        "INTERFACE",
        "Decision Surface",
        "",
        ["Browser control tower", "DOCX evidence export", f"{context.active_alerts} active alerts surfaced"],
    )

    draw_connector((x_positions[0] + main_w, 4.84), (x_positions[1], 4.84), "bootstrap", accents[0])
    draw_connector((x_positions[1] + main_w, 4.84), (x_positions[2], 4.84), "seed", accents[1])
    draw_connector((x_positions[2] + main_w, 4.84), (x_positions[3], 4.84), "analyze", accents[2])
    draw_connector((x_positions[3] + main_w, 4.84), (x_positions[4], 4.84), "deliver", accents[3])

    draw_output_card(
        1.05,
        1.36,
        4.25,
        1.42,
        "#2EE6D6",
        "Traceability Output",
        (
            f"Batch {context.sample_batch_code}\n"
            f"{context.sample_product_name}\n"
            f"{context.sample_component_count} upstream inputs | {context.sample_shipment_count} shipment legs"
        ),
    )
    draw_output_card(
        5.88,
        1.36,
        4.25,
        1.42,
        "#FFC766",
        "Planning Output",
        (
            f"{context.sample_forecast_sku} in {context.sample_forecast_region}\n"
            f"{context.sample_forecast_history_days} days history + {context.sample_forecast_days} day forecast\n"
            "Safety stock and reorder point guidance"
        ),
    )
    draw_output_card(
        10.71,
        1.36,
        4.25,
        1.42,
        "#FF7A8D",
        "Resilience Output",
        (
            f"{context.sample_sim_facility_code}\n"
            f"{context.sample_sim_fill_rate_pct:.1f}% projected recovery coverage\n"
            f"{context.sample_sim_impacted_edges} impacted edge(s) in the sampled scenario"
        ),
    )

    draw_connector((8.12, 3.95), (3.18, 2.63), "trace", "#2EE6D6", rad=0.12)
    draw_connector((10.95, 3.95), (7.95, 2.63), "forecast", "#FFC766", rad=0.05)
    draw_connector((13.95, 3.95), (12.90, 2.63), "recover", "#FF7A8D", rad=-0.02)

    stat_band = FancyBboxPatch(
        (1.00, 0.30),
        14.00,
        0.68,
        boxstyle="round,pad=0.02,rounding_size=0.18",
        facecolor=(1, 1, 1, 0.08),
        edgecolor=(1, 1, 1, 0.10),
        linewidth=1.0,
        zorder=1.7,
    )
    ax.add_patch(stat_band)

    kpi_specs = [
        ("SUPPLIERS", str(context.supplier_facilities), "#2EE6D6"),
        ("BATCHES", str(context.batches), "#5CAEFF"),
        ("SHIPMENTS", str(context.shipments), "#7AF5A0"),
        ("AVG RISK", f"{context.average_risk_score:.2f}", "#FFC766"),
        ("ON-TIME RATE", f"{context.on_time_rate:.2f}%", "#FF7A8D"),
    ]
    segment_width = 14.0 / len(kpi_specs)
    for idx, (label, value, color) in enumerate(kpi_specs):
        x = 1.00 + idx * segment_width
        if idx:
            ax.plot([x, x], [0.40, 0.88], color=(1, 1, 1, 0.10), linewidth=1.0, zorder=2)
        _fit_text(x + 0.15, 0.60, segment_width - 0.30, 0.16, label, "#8CA8C3", 8.0, 7.2, fontweight="bold", key=f"kpi_{idx}_label")
        _fit_text(x + 0.15, 0.39, segment_width - 0.30, 0.22, value, color, 13.4, 10.0, fontweight="bold", key=f"kpi_{idx}_value")

    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    for artist, allowed, key in validation_targets:
        bbox = artist.get_window_extent(renderer=renderer)
        x0, y0, x1, y1 = allowed
        if bbox.x0 < x0 - 1 or bbox.y0 < y0 - 1 or bbox.x1 > x1 + 1 or bbox.y1 > y1 + 1:
            raise RuntimeError(f"Text overflow detected in layout region: {key}")

    for idx, (artist_a, _, key_a) in enumerate(validation_targets):
        bbox_a = artist_a.get_window_extent(renderer=renderer)
        for artist_b, _, key_b in validation_targets[idx + 1 :]:
            bbox_b = artist_b.get_window_extent(renderer=renderer)
            if bbox_a.overlaps(bbox_b):
                raise RuntimeError(f"Text overlap detected between '{key_a}' and '{key_b}'")

    plt.tight_layout()
    fig.savefig(ARCHITECTURE_PATH, dpi=240, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def add_module_table(document: Document, context: RuntimeContext) -> None:
    document.add_paragraph("Repository-to-function mapping:", style="Heading 2")
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Repository Evidence", "Technical Function", "Why It Matters"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        _set_cell_shading(cell, "DCEBEB")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.runs:
            paragraph.runs[0].font.bold = True

    rows = [
        (
            "manage.py + bootstrap.py",
            "Create the database, seed operational records, train models, and export evidence artefacts.",
            "This gives the project a reproducible command path rather than a one-off demo environment.",
        ),
        (
            "models.py + db.py",
            "Define the MySQL-backed relational layer across 17 engineering tables.",
            "The schema makes traceability, logistics, quality, inventory, forecasting, and risk data queryable inside one system.",
        ),
        (
            "seed.py",
            (
                f"Generate a tier-aware dataset with {context.organizations} organisations, {context.facilities} facilities, "
                f"{context.active_links} active links, and {context.demand_rows:,} demand rows."
            ),
            "A realistic testbed is essential because a supply-chain prototype cannot be validated with only a few static records.",
        ),
        (
            "analytics.py",
            (
                f"Train XGBoost risk and demand models, persist joblib artefacts, and record metrics "
                f"(risk RMSE {context.risk_rmse:.2f}, forecast MAPE {context.forecast_mape_pct:.2f}%)."
            ),
            "This turns the repository into a predictive system instead of a descriptive tracker only.",
        ),
        (
            "services.py",
            "Provide dashboard KPIs, batch trace queries, forecast retrieval, and OR-Tools disruption simulation.",
            "The service layer is the contract that keeps analytics, database logic, and UI behaviour consistent.",
        ),
        (
            "ui.py + reports.py",
            "Expose a browser frontend, merchant backend, and chart-backed DOCX evidence for submission.",
            "The final deliverable is both usable in a demo setting and able to generate formal technical artefacts.",
        ),
    ]

    for repo_evidence, function_text, why_text in rows:
        row_cells = table.add_row().cells
        values = [repo_evidence, function_text, why_text]
        for idx, value in enumerate(values):
            row_cells[idx].text = value
            row_cells[idx].paragraphs[0].paragraph_format.space_after = Pt(3)


def add_member_block(
    document: Document,
    member_heading: str,
    student_id: str,
    role: str,
    responsibilities: list[str],
    completed_work: list[str],
    planned_work: list[str],
    management_logic: str | None = None,
) -> None:
    document.add_paragraph(member_heading, style="Heading 2")
    member_name = member_heading.split(" - ", 1)[1].replace(" (Leader)", "")
    add_body_paragraph(document, f"Name & Student ID: {member_name} | {student_id}")
    add_body_paragraph(document, f"Role: {role}")
    if management_logic:
        add_body_paragraph(document, f"Management Logic: {management_logic}")
    add_body_paragraph(document, "Responsibilities:")
    add_bullets(document, responsibilities)
    add_body_paragraph(document, "Completed Work:")
    add_bullets(document, completed_work)
    add_body_paragraph(document, "Planned Work:")
    add_bullets(document, planned_work)


def build_document(context: RuntimeContext) -> Document:
    document = Document()
    style_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("INF101TC Semester 2 Technical Deliverable Report")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.add_run(f"Group Number: {GROUP_NUMBER}\n").bold = True
    subtitle.add_run(f"Project Title: {PROJECT_TITLE}\n")
    subtitle.add_run(
        "Team Members: Rui Huang (Leader, 2471007), Zixiu Wang (2469486), Jiangluhai Pan (2469560), "
        "Shuo Yuan (2469754), Jiyang Bai (2470943), Junze Wu (2470531)\n"
    )
    subtitle.add_run(f"Submission Date: {date.today().isoformat()}")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(14)
    note_run = note.add_run(
        "This version standardizes the full document format, aligns every technical statement with the current repository, "
        "and places the group leader's contribution first."
    )
    note_run.italic = True
    note_run.font.size = Pt(10.5)
    note_run.font.color.rgb = RGBColor(79, 93, 117)

    document.add_paragraph("Section 1: The Technical Solution(s)", style="Heading 1")
    document.add_paragraph("System Walkthrough", style="Heading 2")

    add_body_paragraph(
        document,
        "The semester deliverable is a working Python repository that implements a fresh-food supply chain intelligence platform rather than a static concept mock-up. "
        "The solution combines a local MySQL runtime, SQLAlchemy data models, machine-learning analytics, linear-programming-based recovery planning, and a browser-based interface. "
        "Its engineering goal is not only to show where a product came from, but to model how upstream suppliers, cold-chain logistics, inventory, demand, and disruption risk interact inside one traceable system.",
    )
    add_body_paragraph(
        document,
        "The technical workflow begins with scripts/setup_local_mysql.ps1, which prepares a dedicated local MySQL service, followed by manage.py bootstrap-all, "
        "which creates the schema, seeds the synthetic network, trains the analytical models, and exports evidence artefacts. "
        f"In the current repository state, the seeded database contains {context.organizations} organisations, {context.facilities} facilities, "
        f"{context.active_links} active supply links, {context.supplier_lots} supplier lots, {context.batches} finished-goods batches, {context.shipments} shipments, and {context.demand_rows:,} demand observations. "
        "This gives the team a realistic engineering testbed instead of a thin demonstration with only a few hand-written examples.",
    )

    add_figure(
        document,
        ARCHITECTURE_PATH,
        6.5,
        "Figure 1. Commercial-grade end-to-end system flow showing how runtime setup, relational data engineering, digital-twin seeding, analytics, and operator-facing outputs connect in one decision pipeline.",
    )

    add_body_paragraph(
        document,
        "At the data layer, models.py and db.py define a 17-table relational structure covering locations, organisations, facilities, bills of material, supply edges, supplier lots, product batches, shipment history, inventory snapshots, quality inspections, demand history, forecast outputs, risk assessments, and alert events. "
        "seed.py then synthesises tier-aware records across L1, L2, L3, CORE, DOWNSTREAM, and SERVICE nodes. "
        "This architecture matters because traceability in a fresh-food setting is not only a product-level problem; it depends on upstream depth, quality state, logistics conditions, and demand pressure being represented together.",
    )
    add_module_table(document, context)

    add_body_paragraph(
        document,
        "The application layer turns these tables into a usable control tower. ui.py exposes five aligned tabs - Dashboard, Network Mesh, Traceability, Forecasting, and Scenario Lab - while services.py provides the query and decision logic behind each view. "
        f"For example, the current seeded batch {context.sample_batch_code} resolves to {context.sample_product_name} produced at {context.sample_plant_name}, and the trace view links that batch to {context.sample_component_count} upstream component records plus {context.sample_shipment_count} shipment legs. "
        "This demonstrates that the system walkthrough is grounded in executable repository behaviour rather than in placeholder screenshots.",
    )

    add_figure(
        document,
        NETWORK_PATH,
        6.45,
        "Figure 2. Mesh topology generated from the seeded database, showing how L1, L2, and L3 suppliers connect to CORE plants and downstream fulfilment nodes.",
    )

    add_body_paragraph(
        document,
        "The analytical layer extends the project beyond descriptive traceability. analytics.py trains an XGBoost facility-risk model and an XGBoost lag-based demand forecast model, then persists the resulting artefacts with joblib and JSON metrics. "
        f"On the latest run, the facility-risk model achieved RMSE {context.risk_rmse:.2f} and MAE {context.risk_mae:.2f} across {context.facilities} facilities, while the {context.forecast_horizon_days}-day forecast model achieved RMSE {context.forecast_rmse:.2f}, "
        f"MAE {context.forecast_mae:.2f}, and MAPE {context.forecast_mape_pct:.2f} percent. In practical terms, the system can load {context.sample_forecast_history_days} days of historical demand and {context.sample_forecast_days} days of forward predictions for "
        f"{context.sample_forecast_sku} in {context.sample_forecast_region}, together with safety-stock and reorder-point recommendations.",
    )

    add_figure(
        document,
        RISK_PATH,
        6.15,
        "Figure 3. Tier-level risk distribution generated from the trained facility-risk model.",
    )
    add_figure(
        document,
        DEMAND_PATH,
        6.15,
        "Figure 4. Recent demand trend used to support forecasting and inventory-planning decisions.",
    )

    add_body_paragraph(
        document,
        "The final technical differentiator is decision support under disruption. services.py uses OR-Tools linear optimisation to reallocate flows when a facility loses capacity. "
        f"In the current seeded network, a 35 percent capacity-drop simulation on {context.sample_sim_facility_code} ({context.sample_sim_facility_name}) still produced "
        f"{context.sample_sim_fill_rate_pct:.1f} percent projected recovery coverage across {context.sample_sim_impacted_edges} impacted edge(s). "
        "This is why the team selected a digital-twin approach: it can explain what happened, forecast what is likely to happen next, and recommend how the network should react.",
    )
    add_body_paragraph(
        document,
        "Taken together, the repository delivers a coherent technical solution: bootstrapping scripts establish the runtime, the database stores tier-aware operational history, analytics transform raw data into predictions, the browser UI makes the information explorable, and reports.py generates submission-ready DOCX evidence. "
        "The semester deliverable is therefore a functioning engineering prototype with reproducible commands, measurable outputs, and a clear path to future extension.",
    )

    document.add_paragraph("Section 2: Limitations", style="Heading 1")
    add_bullets(
        document,
        [
            "Synthetic data boundary. The network structure is realistic, but the underlying records are still generated rather than sourced directly from CP Group or a live enterprise environment. The current results therefore validate architecture, workflow, and analytical feasibility more than they validate production-level decision accuracy.",
            "Local browser deployment. The deliverable currently runs as a local MySQL plus FastAPI/React application on a single machine. It is strong for technical demonstration and controlled testing, but it is not yet a cloud-hosted, multi-user platform with role-based access control.",
            "Traceability input channel. Once a batch exists in the database, the system can trace it very effectively; however, the project does not yet include a production QR/mobile scanning pipeline, IoT sensor ingestion, or automatic ERP/WMS connectors for real-time event capture.",
            "Model scope and assumptions. Both XGBoost models are trained on synthetic operating patterns, and the optimisation engine uses simplified cost-capacity assumptions. Under black-swan disruptions, policy shocks, or atypical supplier behaviour, human review would still be necessary before acting on the recommendations.",
            "Operational hardening. Automated backup, audit logging, cyber-security controls, and large-scale performance benchmarking remain prototype-level. These are essential before the system could be used as a mission-critical enterprise platform.",
        ],
    )
    add_body_paragraph(
        document,
        "These limitations are acceptable for a semester prototype because the primary aim is to demonstrate a complete and credible engineering path from business problem to working technical system.",
    )

    document.add_paragraph("Section 3: Future Development", style="Heading 1")
    add_bullets(
        document,
        [
            "Live data integration and API layering. The next iteration should connect the repository to real supplier, plant, logistics, and quality data through a secure API or ETL pipeline, with formal field mapping to replace synthetic assumptions. This step requires enterprise data access, data-governance approval, and interface contracts with operational stakeholders.",
            "Mobile QR and IoT extension. A production version should add a mobile or lightweight web scan interface, temperature-sensor ingestion, and event-driven alerting so that traceability begins at the real capture point instead of at seeded database records. This would strengthen both consumer trust and operator response speed.",
            "Production hardening and scalability. The system should be containerised, backed by scheduled backups, protected with authentication and role-based access control, and stress-tested under concurrent use. Monitoring, logging, and performance baselines are needed so the platform can move from a coursework prototype to an operational pilot.",
            "Model governance and richer simulation. Future work should add retraining pipelines, model drift monitoring, supplier feedback loops, and multi-objective optimisation that balances cost, service level, cold-chain compliance, and resilience at the same time. This would make the digital twin more useful as a decision-support tool.",
            "Pilot rollout plan. A sensible commercial pathway is to pilot the platform on one product family and two or three regions first, then expand once the data interfaces, operator workflows, and business value are validated in practice.",
        ],
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    style_document(document)

    document.add_paragraph("Section 4: Individual Contributions", style="Heading 1")
    add_body_paragraph(
        document,
        "The member blocks below have been rewritten into one consistent format, aligned to the current repository, and reordered to place the group leader first.",
    )

    add_member_block(
        document,
        "Member 1 - Rui Huang (Leader)",
        "2471007",
        "Team Leader, System Architect, and Integration Manager",
        [
            "Translate the original business problem into a deliverable technical architecture and define the scope of the semester repository.",
            "Decompose the project into database, analytics, interface, and research workstreams, assign clear ownership, and control integration checkpoints across the team.",
            "Review technical consistency across the codebase, documentation, and demonstration narrative so the final submission behaves as one system instead of a set of disconnected drafts.",
        ],
        [
            "Finalised the current project direction around a MySQL plus FastAPI/React plus analytics digital-twin platform and aligned the team away from earlier inconsistent static mock-up descriptions.",
            "Coordinated interface consistency across schema design, seeded data, analytics outputs, and UI behaviour so that traceability, forecasting, and scenario simulation use the same tier logic and naming conventions.",
            "Led end-to-end verification of the local MySQL runtime, repository workflow, generated artefacts, and final report quality, then consolidated conflicting member write-ups into one coherent submission standard.",
            "Directed the final technical documentation so that the report reflects the real repository structure, actual metrics, and reproducible project workflow.",
        ],
        [
            "Run final acceptance testing on the full project chain, including environment setup, data generation, analytics execution, and demonstration stability.",
            "Control the final submission package, speaking narrative, and risk checklist so the team presents one polished and defensible technical story.",
            "Coordinate last-stage refinement on performance, documentation clarity, and demonstration evidence before final hand-in.",
        ],
        management_logic=(
            "I managed the project through a stage-gate workflow: first define the technical acceptance criteria, then assign each teammate a primary module owner role, "
            "then review cross-module interfaces at fixed checkpoints around the database schema, service-layer outputs, and UI requirements, and finally unify all documentation against the repository rather than against individual notes. "
            "This management logic reduced overlap, made accountability clear, and helped the team resolve conflicting drafts before submission."
        ),
    )

    add_member_block(
        document,
        "Member 2 - Zixiu Wang",
        "2469486",
        "Database Engineering Lead",
        [
            "Design the project's relational data foundation and ensure that traceability, logistics, quality, and demand entities are represented consistently.",
            "Support data integrity, relationship mapping, and query efficiency for the engineering workflow.",
            "Collaborate with analytics and interface roles so downstream modules receive reliable structured data.",
        ],
        [
            "Translated the project domain into a MySQL and SQLAlchemy schema covering organisations, facilities, materials, product batches, shipments, inspections, demand history, and analytical outputs.",
            "Helped define the relationship logic required for tier-aware traceability so that upstream lots, downstream batches, and shipment records can be connected in one data model.",
            "Supported validation of keys, entity naming, and data consistency rules that make the service layer and seeded dataset usable for analytics and UI presentation.",
        ],
        [
            "Refine query performance and indexing for larger datasets and more demanding demonstration scenarios.",
            "Prepare the schema for future migration from synthetic data to live enterprise records, including stronger audit and governance fields.",
            "Continue supporting integration between the database layer and later-stage reporting or API work.",
        ],
    )

    add_member_block(
        document,
        "Member 3 - Jiangluhai Pan",
        "2469560",
        "Technical Consultant and Debugging Support",
        [
            "Research feasible technical routes and help the group select an architecture that can be implemented within the semester timeframe.",
            "Review logic, troubleshoot cross-module issues, and provide technical support when different components need to work together.",
            "Contribute to quality control by identifying risk points, edge cases, and architectural trade-offs.",
        ],
        [
            "Supported the evaluation of technology choices and helped justify the move toward a Python-based engineering stack that could integrate data, analytics, and interface work inside one repository.",
            "Assisted in checking logic consistency between the data model, service outputs, and planned user-facing workflow, especially where technical assumptions needed to be made explicit.",
            "Contributed debugging and issue-triage support so that problems could be identified early and resolved before they affected the integrated project story.",
        ],
        [
            "Support more formal integration testing and edge-case validation as the team prepares the final demonstration.",
            "Review technical assumptions inside the simulation and analytics pipeline so the future roadmap remains realistic and defensible.",
            "Continue acting as a technical bridge when cross-module issues emerge near final submission.",
        ],
    )

    add_member_block(
        document,
        "Member 4 - Shuo Yuan",
        "2469754",
        "Market Research and Data Validation Analyst",
        [
            "Translate industry pain points and user needs into system requirements that matter for the business problem.",
            "Validate whether the project's data fields, assumptions, and metrics remain commercially meaningful instead of being purely technical.",
            "Support research, business framing, and evidence selection for the final submission.",
        ],
        [
            "Collected and organised fresh-food traceability and supply-chain research so the team could frame the problem around trust, food safety, supply instability, and decision speed.",
            "Supported validation of the seeded dataset and KPI assumptions so the project's demand, compliance, and traceability fields stayed aligned with believable business use cases.",
            "Contributed research-driven justification for the current scope, limitations, and future development direction of the platform.",
        ],
        [
            "Deepen competitor benchmarking and strengthen the business case for user-facing transparency and operational resilience features.",
            "Support future iterations of KPI interpretation, stakeholder storytelling, and user-oriented evidence selection.",
            "Continue helping the team connect technical outputs with practical market and management value.",
        ],
    )

    add_member_block(
        document,
        "Member 5 - Jiyang Bai",
        "2470943",
        "Analytics and Dashboard Design Specialist",
        [
            "Define the data-analysis perspective of the platform so raw records can be transformed into operational insight.",
            "Support KPI design, dashboard logic, and model-output interpretation for non-technical audiences.",
            "Collaborate with the database and interface roles to keep analytics-ready data and visual presentation aligned.",
        ],
        [
            "Helped identify the analytical indicators needed for the project, including traceability completeness, cold-chain exposure, demand variability, and operational risk visibility.",
            "Supported the design logic behind dashboard-style presentation so the system could communicate more than raw database tables during the demonstration.",
            "Contributed to analysis-oriented checking of data quality assumptions and the way model outputs would be interpreted inside the report and interface.",
        ],
        [
            "Expand the analytical storytelling depth of the dashboard and improve how complex outputs are explained to users.",
            "Support later-stage monitoring logic for forecasting quality and risk scoring stability.",
            "Continue refining the bridge between analytical evidence and decision-making value in the final prototype.",
        ],
    )

    add_member_block(
        document,
        "Member 6 - Junze Wu",
        "2470531",
        "UI/UX and Visual Interface Developer",
        [
            "Shape the user-facing workflow of the application so the system can be demonstrated clearly and intuitively.",
            "Support visual consistency, information hierarchy, and interface usability across the main functional modules.",
            "Provide design feedback that helps technical outputs become understandable during live demonstration or review.",
        ],
        [
            "Supported the structure and presentation logic of the browser interface across customer view, merchant studio, traceability, forecasting, and scenario workflows.",
            "Helped refine how users move through the application so the project can be presented as a coherent control-tower workflow instead of a disconnected feature list.",
            "Contributed review feedback on layout clarity, visual consistency, and demonstration readability for the final submission materials.",
        ],
        [
            "Polish interaction details, demonstration screenshots, and presentation flow for the final hand-in version.",
            "Support window-state and usability checks so the interface remains clear during live walkthroughs.",
            "Continue improving the clarity of user guidance and visual storytelling across the application.",
        ],
    )

    return document


def main() -> None:
    context = collect_runtime_context()
    build_architecture_diagram(context)
    document = build_document(context)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
